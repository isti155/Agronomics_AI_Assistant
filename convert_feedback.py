"""
Convert farmer_feedback_questionnaire.html to PDF and DOCX.
"""
import os, re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HTML_PATH = os.path.join(os.path.dirname(__file__), '..', 'farmer_feedback_questionnaire.html')
OUT_DIR = os.path.join(os.path.dirname(__file__), '..') 

def clean(text):
    return (text or '').strip().replace('\xa0', ' ')

def add_styled_paragraph(doc, text, bold=False, italic=False, size=11, color=None, space_after=6, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p

def process_page(doc, page_div, page_title):
    """Process a single tab/page from the HTML into the DOCX."""
    doc_pages = page_div.find_all('div', class_='doc-page')
    
    for dp in doc_pages:
        # Cover page
        cover = dp.find('div', class_='cover')
        if cover:
            label = cover.find('p', class_='cover-label')
            if label:
                add_styled_paragraph(doc, clean(label.text), bold=True, size=16, space_after=12,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        # Participant header
        ph = dp.find('p', class_='participant-header')
        if ph:
            text = clean(ph.get_text(separator='\n'))
            add_styled_paragraph(doc, text, italic=True, size=9, color=(85,85,85), space_after=14)

        # Process all children in order
        for el in dp.children:
            if not hasattr(el, 'name') or el.name is None:
                continue
            tag = el.name
            cls = el.get('class', [])

            if tag == 'h1':
                add_styled_paragraph(doc, clean(el.text), bold=True, size=16, space_after=10)
            elif tag == 'h2':
                add_styled_paragraph(doc, clean(el.text), bold=True, size=13, space_after=8)
            elif tag == 'h3':
                add_styled_paragraph(doc, clean(el.text), bold=True, size=11, space_after=6)
            elif tag == 'p':
                if 'participant-header' in cls:
                    continue
                if 'note' in cls:
                    add_styled_paragraph(doc, clean(el.text), italic=True, size=10, color=(68,68,68))
                else:
                    txt = clean(el.get_text())
                    is_bold = el.find('strong') is not None
                    add_styled_paragraph(doc, txt, bold=is_bold, size=11)
            elif tag == 'hr':
                add_styled_paragraph(doc, '—' * 50, size=8, color=(200,200,200), space_after=10)
            elif tag == 'ol':
                for i, li in enumerate(el.find_all('li', recursive=False), 1):
                    li_text = []
                    for child in li.children:
                        if hasattr(child, 'name') and child.name in ('ul', 'div', 'table'):
                            continue
                        li_text.append(clean(child.string or child.get_text() if hasattr(child, 'get_text') else str(child)))
                    text = ' '.join(t for t in li_text if t)
                    add_styled_paragraph(doc, f"{i}. {text}", size=11, space_after=4)
                    # Sub-options
                    sub_ul = li.find('ul', class_='options')
                    if sub_ul:
                        for opt in sub_ul.find_all('li'):
                            add_styled_paragraph(doc, f"    {clean(opt.text)}", size=10, space_after=2)
                    # Scale row
                    sr = li.find('div', class_='scale-row')
                    if sr:
                        vals = [clean(d.text) for d in sr.find_all('div')]
                        sels = ['selected' in (d.get('class') or []) for d in sr.find_all('div')]
                        row_text = '  '.join(f"[{v}]" if s else f" {v} " for v, s in zip(vals, sels))
                        add_styled_paragraph(doc, f"    Scale: {row_text}", size=10, space_after=4)
            elif tag == 'ul' and 'options' in cls:
                for li in el.find_all('li'):
                    add_styled_paragraph(doc, f"    {clean(li.text)}", size=10, space_after=2)
            elif tag == 'table' and 'sus-table' in cls:
                rows = el.find_all('tr')
                if not rows:
                    continue
                headers = [clean(th.text) for th in rows[0].find_all(['th', 'td'])]
                ncols = len(headers)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = 'Table Grid'
                for ci, h in enumerate(headers):
                    cell = table.rows[0].cells[ci]
                    cell.text = h
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(9)
                for ri, row in enumerate(rows[1:], 1):
                    cells = row.find_all('td')
                    for ci, td in enumerate(cells):
                        if ci < ncols:
                            table.rows[ri].cells[ci].text = clean(td.text)
                            for p in table.rows[ri].cells[ci].paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(9)
                doc.add_paragraph()  # spacing
            elif tag == 'div':
                if 'cb-item' in cls:
                    cb = el.find('span', class_='cb-box')
                    checked = cb and 'checked' in (cb.get('class') or [])
                    mark = '☑' if checked else '☐'
                    label = clean(el.get_text())
                    add_styled_paragraph(doc, f"  {mark}  {label}", size=10, space_after=2)
                elif 'followup' in cls:
                    add_styled_paragraph(doc, '', size=6, space_after=2)
                    for child in el.children:
                        if not hasattr(child, 'name') or child.name is None:
                            continue
                        if child.name == 'p':
                            ccls = child.get('class', [])
                            if 'fq' in ccls:
                                add_styled_paragraph(doc, clean(child.text), italic=True, size=10)
                            elif 'ans' in ccls:
                                add_styled_paragraph(doc, clean(child.text), bold=True, size=10)
                            else:
                                add_styled_paragraph(doc, clean(child.text), size=10)

def html_to_docx(html_path, output_path):
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'lxml')
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Agronomics – User Feedback Questionnaire', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    pages = soup.find_all('div', class_='page')
    tab_names = ['Survey Format', 'id_1 – Mr. Rahim', 'id_2 – Ms. Fatema', 
                 'id_3 – Mr. Karim', 'id_4 – Ms. Rekha', 'id_5 – Mr. Jahangir']
    
    for i, page in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        name = tab_names[i] if i < len(tab_names) else f'Tab {i+1}'
        process_page(doc, page, name)
    
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

def generate_print_html(html_path):
    """Create a print-ready HTML with all pages visible for PDF generation."""
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    # Make all pages visible and hide tab bar
    content = content.replace('display: none;', 'display: block;')
    content = content.replace('.page.active {\n      display: block;\n    }', '')
    # Hide tab bar
    content = content.replace('<div class="tab-bar">', '<div class="tab-bar" style="display:none !important;">')
    return content

def html_to_pdf(html_path, output_path):
    """Generate PDF using browser print. Creates a print-ready HTML first."""
    print_html = generate_print_html(html_path)
    print_html_path = output_path.replace('.pdf', '_print.html')
    with open(print_html_path, 'w', encoding='utf-8') as f:
        f.write(print_html)
    
    # Try multiple PDF generation methods
    pdf_generated = False
    
    # Method 1: Try pdfkit/wkhtmltopdf
    try:
        import pdfkit
        pdfkit.from_file(print_html_path, output_path)
        pdf_generated = True
        print(f"PDF saved (pdfkit): {output_path}")
    except Exception as e:
        print(f"pdfkit not available: {e}")
    
    # Method 2: Try weasyprint
    if not pdf_generated:
        try:
            from weasyprint import HTML
            HTML(filename=print_html_path).write_pdf(output_path)
            pdf_generated = True
            print(f"PDF saved (weasyprint): {output_path}")
        except Exception as e:
            print(f"weasyprint not available: {e}")
    
    # Method 3: Try Chrome/Edge headless
    if not pdf_generated:
        import subprocess, shutil
        browser_paths = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        ]
        browser = next((p for p in browser_paths if os.path.exists(p)), None)
        if browser:
            try:
                abs_print_html = os.path.abspath(print_html_path)
                subprocess.run([
                    browser, '--headless', '--disable-gpu',
                    f'--print-to-pdf={os.path.abspath(output_path)}',
                    '--no-margins', f'file:///{abs_print_html}'
                ], timeout=30, capture_output=True)
                if os.path.exists(output_path):
                    pdf_generated = True
                    print(f"PDF saved (Chrome/Edge headless): {output_path}")
            except Exception as e:
                print(f"Chrome/Edge headless failed: {e}")

    if not pdf_generated:
        print(f"\nPDF could not be auto-generated (no PDF engine found).")
        print(f"A print-ready HTML has been saved at: {print_html_path}")
        print(f"To create the PDF manually:")
        print(f"  1. Open '{print_html_path}' in Chrome/Edge")
        print(f"  2. Press Ctrl+P")
        print(f"  3. Choose 'Save as PDF' and save")

    return pdf_generated, print_html_path

if __name__ == '__main__':
    html_path = os.path.abspath(HTML_PATH)
    docx_out = os.path.join(os.path.abspath(OUT_DIR), 'farmer_feedback_questionnaire.docx')
    pdf_out = os.path.join(os.path.abspath(OUT_DIR), 'farmer_feedback_questionnaire.pdf')
    
    print("=" * 50)
    print("Converting farmer feedback questionnaire...")
    print("=" * 50)
    
    # Generate DOCX
    print("\n[1/2] Generating DOCX...")
    html_to_docx(html_path, docx_out)
    
    # Generate PDF
    print("\n[2/2] Generating PDF...")
    html_to_pdf(html_path, pdf_out)
    
    print("\n" + "=" * 50)
    print("Done!")
